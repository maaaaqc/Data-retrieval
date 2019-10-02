import json
from docx import Document
from docx.shared import Pt
from pathlib import Path

CONFIG = Path.cwd() / "config.json"

def convert_word(dictionary, word):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.2
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    with open(CONFIG) as json_file:
        target = json.load(json_file)
    for key in target:
        head = document.add_paragraph()
        head.add_run(target[key].upper()).bold=True
        if not key == "nlpsinfo" \
            and not key == "profile" \
            and not "ccss" in key:
            for i in dictionary[key]:
                document.add_paragraph(i)
        else:
            table = document.add_table(rows=len(dictionary[key]), cols=len(dictionary[key][0]))
            table.style = 'TableGrid'
            row = 0
            for val in dictionary[key]:
                col = 0
                for i in val:
                    table.rows[row].cells[col].text = i
                    col += 1
                row += 1
        document.add_paragraph()
    document.save(word)